import streamlit as st
import anthropic
import os
import json
import io
import re
from datetime import date
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    import fitz
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

st.set_page_config(
    page_title="ADOSH Risk Bot - ENGC",
    page_icon="🦺",
    layout="wide"
)

st.title("🦺 ADOSH Risk Assessment Bot")
st.caption("Exeed National General Contracting LLC | Bloom Living Projects | Abu Dhabi")
st.markdown("---")

# ── ADOSH CoP REFERENCE MAP ──
ADOSH_COP_MAP = """
ADOSH ABU DHABI CODES OF PRACTICE (CoP) REFERENCE GUIDE:
Always cite the correct CoP number in your legal references based on the activity.

CoP 1.0 - Hazardous Materials
CoP 1.1 - Management of Asbestos Containing Materials
CoP 1.2 - Lead Exposure Management
CoP 2.0 - Personal Protective Equipment (PPE)
CoP 3.0 - Occupational Noise
CoP 3.1 - Vibration
CoP 4.0 - First Aid and Medical Emergency Treatment
CoP 5.0 - Occupational Health Screening and Medical Surveillance
CoP 8.0 - General Workplace Amenities
CoP 9.0 - Workplace Wellness
CoP 9.1 - New and Expectant Mothers
CoP 9.2 - Managing Work-Related Stress
CoP 10.0 - Rehabilitation and Return to Work
CoP 11.0 - Safety in the Heat (MOHRE Resolution 44/2022 - midday ban 12:30-15:00, June 15 to September 15)
CoP 12.0 - Prevention and Control of Legionnaires Disease
CoP 13.0 - Violence in the Workplace
CoP 14.0 - Manual Handling and Ergonomics
CoP 14.1 - Manual Tasks Involving the Handling of People
CoP 15.0 - Electrical Safety
CoP 16.0 - OSH Requirements for People with Special Needs
CoP 17.0 - Safety Signage and Signals
CoP 18.0 - Employer Supplied Accommodation General Requirements
CoP 18.1 - Temporary Employer Supplied Accommodation
CoP 19.0 - Occupational Food Handling and Food Preparation Areas
CoP 20.0 - Safety in Design (Construction)
CoP 21.0 - Permit to Work Systems
CoP 22.0 - Barricading of Hazards
CoP 23.0 - Working at Heights
CoP 24.0 - Lock-out Tag-out (Isolation)
CoP 25.0 - Driver Fatigue Prevention
CoP 26.0 - Scaffolding
CoP 27.0 - Confined Spaces
CoP 28.0 - Hot Work Operations
CoP 29.0 - Excavation Work
CoP 30.0 - Lone Working and/or in Remote Locations
CoP 31.0 - Working on, Over or Adjacent to Water
CoP 33.0 - Working On or Adjacent to a Road
CoP 33.1 - Traffic Incident Site Management
CoP 34.0 - Safe Use of Lifting Equipment and Lifting Accessories
CoP 35.0 - Portable Power Tools
CoP 36.0 - Plant and Equipment
CoP 37.0 - Ladders
CoP 38.0 - Concrete Placing Equipment
CoP 39.0 - Overhead and Underground Services
CoP 40.0 - False Work (Formwork)
CoP 41.0 - Steel Erection
CoP 42.0 - Pre-Cast Construction
CoP 43.0 - Temporary Structures
CoP 44.0 - Traffic Management and Logistics
CoP 45.0 - Underwater Activities
CoP 46.0 - Underground Construction
CoP 46.1 - Construction of Water Wells
CoP 47.0 - Machine Guarding
CoP 48.0 - Spray Finishing
CoP 49.0 - Compressed Gases and Air
CoP 50.0 - Abrasive Blasting and Associated Protective Coating Work
CoP 51.0 - Powered Lift Trucks
CoP 52.0 - Local Exhaust Ventilation
CoP 53.0 - OSH Management During Construction Work
CoP 53.1 - OSH Construction Management Plan
CoP 54.0 - Waste Management

ACTIVITY TO CoP MAPPING (always use these):
- Working at Heights → CoP 23.0, CoP 26.0 (Scaffolding), CoP 37.0 (Ladders), CoP 2.0 (PPE)
- Scaffolding → CoP 26.0, CoP 23.0, CoP 2.0, CoP 17.0
- Excavation → CoP 29.0, CoP 39.0, CoP 22.0, CoP 2.0, CoP 21.0
- Lifting Operations → CoP 34.0, CoP 36.0, CoP 21.0, CoP 2.0, CoP 44.0
- Crane Operations → CoP 34.0, CoP 36.0, CoP 44.0, CoP 21.0
- Electrical Works → CoP 15.0, CoP 24.0, CoP 21.0, CoP 2.0
- Confined Space → CoP 27.0, CoP 21.0, CoP 4.0, CoP 2.0
- Hot Work / Welding → CoP 28.0, CoP 21.0, CoP 2.0, CoP 49.0
- Chemical Handling → CoP 1.0, CoP 2.0, CoP 52.0, CoP 54.0
- Heat Stress → CoP 11.0, MOHRE Resolution 44/2022
- Fire Prevention → CoP 21.0, CoP 17.0, CoP 4.0, CoP 2.0
- Concrete Works → CoP 38.0, CoP 40.0, CoP 2.0, CoP 22.0
- Pre-Cast Construction → CoP 42.0, CoP 34.0, CoP 23.0
- Formwork / False Work → CoP 40.0, CoP 23.0, CoP 2.0
- Steel Erection → CoP 41.0, CoP 23.0, CoP 34.0, CoP 2.0
- Manual Handling → CoP 14.0, CoP 2.0, CoP 3.1
- Noise → CoP 3.0, CoP 2.0
- Traffic / Road Works → CoP 33.0, CoP 44.0, CoP 17.0
- Plant and Equipment → CoP 36.0, CoP 35.0, CoP 24.0
- Portable Power Tools → CoP 35.0, CoP 15.0, CoP 2.0
- Permit to Work → CoP 21.0
- PPE → CoP 2.0
- Safety Signage → CoP 17.0
- Waste Management → CoP 54.0
- Night Shift Work → CoP 8.0, CoP 11.0, CoP 4.0
- Underground Services → CoP 39.0, CoP 29.0, CoP 21.0
- Overhead Services → CoP 39.0, CoP 15.0, CoP 21.0
- Compressed Gas → CoP 49.0, CoP 2.0, CoP 21.0
- Abrasive Blasting → CoP 50.0, CoP 2.0, CoP 52.0
- Powered Lift Trucks → CoP 51.0, CoP 36.0, CoP 44.0
- Ladders → CoP 37.0, CoP 23.0, CoP 2.0
- Machine Guarding → CoP 47.0, CoP 36.0, CoP 24.0
- General Construction → CoP 53.0, CoP 53.1, CoP 20.0
"""

SYSTEM_PROMPT = """You are ADOSH Risk Bot for ENGC Abu Dhabi UAE.
Generate Risk Assessments following ADOSH-SF Version 4.0.

""" + ADOSH_COP_MAP + """

COLUMN ORDER IN JSON - CRITICAL - FOLLOW EXACTLY:
The correct column sequence is:
1. sn
2. activity_element
3. hazards
4. who_harmed  ← WHO MIGHT BE HARMED comes BEFORE consequences
5. consequences ← RISK AND CONSEQUENCES comes AFTER who_harmed
6. prob_initial, sev_initial, risk_initial, risk_level_initial
7. controls
8. prob_residual, sev_residual, risk_residual, risk_level_residual

Risk Matrix: Probability x Severity. LOW=1-4, MODERATE=5-9, HIGH=10-14, EXTREME=15-25.

Generate EXACTLY 10 rows specific to the activity requested.
Analyze the activity and break it into 10 real construction steps.
Each row must be a DIFFERENT specific step of that exact activity.
Think like this for any activity:
- Excavation: site setup, marking, breaking, shoring, dewatering, soil removal, underground services, backfill, compaction, reinstatement
- Scaffolding: delivery, ground prep, base plates, standards, ledgers, transoms, boarding, handrails, ladders, inspection
- Electrical: isolation, permit to work, cable routing, pulling, termination, testing, energizing, earthing, panel work, inspection
- Concrete: formwork, reinforcement, delivery, pouring, compaction, curing, stripping, finishing, testing, waste
- Working at Heights: equipment check, edge protection, harness, anchors, platform, tool tethering, weather, rescue plan, supervision, descent
- Fire Prevention: site assessment, ignition sources, storage, detection, suppression, signage, evacuation, training, hot works, inspection
- Chemical Handling: MSDS review, storage, PPE selection, mixing, spillage, disposal, ventilation, emergency, first aid, inspection
- Lifting Operations: pre-lift plan, crane inspection, outrigger setup, rigging, slinging, trial lift, main lift, slewing, load landing, post-lift
- Confined Space: risk assessment, permit to work, atmospheric testing, ventilation, entry, monitoring, communication, rescue, exit, debrief
Apply same thinking to ANY other activity requested.
Row 9 always = Heat Stress (CoP 11.0 / MOHRE Resolution 44/2022, midday ban 12:30-15:00 June 15 to September 15).
Row 10 always = Emergency Response and First Aid (CoP 4.0).
Keep each text field under 50 words. Residual risk must be LOW or MODERATE.
Always include specific ADOSH CoP numbers in the legal_references field.
Always include specific ADOSH CoP numbers inside each control measure where relevant.

OUTPUT RULES - CRITICAL:
1. Output ONLY a JSON object
2. Start with { and end with }
3. Zero text before or after the JSON
4. No markdown, no code blocks, no explanation

JSON format - COLUMN ORDER IS FIXED:
{
  "activity": "name",
  "project": "name",
  "rows": [
    {
      "sn": 1,
      "activity_element": "specific step",
      "hazards": "hazard identified",
      "who_harmed": "who is harmed and exactly how",
      "consequences": "risk and potential consequences",
      "prob_initial": 4,
      "sev_initial": 4,
      "risk_initial": 16,
      "risk_level_initial": "HIGH",
      "controls": "1. Elimination (CoP ref)\n2. Substitution\n3. Engineering control\n4. Administrative (CoP ref)\n5. PPE per CoP 2.0",
      "prob_residual": 2,
      "sev_residual": 3,
      "risk_residual": 6,
      "risk_level_residual": "MODERATE"
    }
  ],
  "legal_references": "CoP 23.0 Working at Heights, CoP 2.0 PPE, CoP 21.0 Permit to Work, ADOSH-SF v4.0"
}"""


# ── CHATBOT SYSTEM PROMPT ──
CHAT_SYSTEM_PROMPT = """You are an expert HSE Assistant for Exeed National
General Contracting LLC (ENGC), Abu Dhabi, UAE.

You specialize in:
- ADOSH-SF Version 4.0 regulations
- Abu Dhabi Codes of Practice (CoP 2.0, 11.0, 15.0, 21.0, 23.0, 26.0, 27.0, 29.0, 34.0, 53.1)
- MOHRE Resolution No. 44/2022 (midday ban 12:30-15:00, June 15 to September 15)
- ISO 45001:2018 and ISO 14001:2015
- Construction site safety in Abu Dhabi
- Heat stress management (CoP 11.0)
- Permit to Work systems (CoP 21.0)
- Working at Heights (CoP 23.0)
- Scaffolding (CoP 26.0)
- Excavation (CoP 29.0)
- Lifting Operations (CoP 34.0)
- Confined Spaces (CoP 27.0)
- Electrical Safety (CoP 15.0)

Always cite the correct ADOSH CoP number in your answers.
Give practical site-specific advice for Abu Dhabi construction projects.
Be professional but easy to understand."""


def get_risk_color(level):
    colors = {
        "LOW": "00AF50",
        "MODERATE": "FFFF00",
        "HIGH": "FFC000",
        "EXTREME": "FF0000"
    }
    return colors.get(str(level).upper(), "FFFFFF")


def shade_cell(cell, hex_color):
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)
    except Exception:
        pass


def extract_json(text):
    text = text.strip()
    if "```json" in text:
        text = text.split("```json")[1]
        if "```" in text:
            text = text.split("```")[0]
        text = text.strip()
    elif "```" in text:
        text = text.split("```")[1]
        if "```" in text:
            text = text.split("```")[0]
        text = text.strip()
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1:
        raise ValueError("No JSON found")
    text = text[start:end + 1]
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    last = text.rfind('"}')
    if last > 0:
        truncated = text[:last + 2]
        ob = truncated.count('{') - truncated.count('}')
        ob2 = truncated.count('[') - truncated.count(']')
        if ob2 > 0:
            truncated += ']' * ob2
        if ob > 0:
            truncated += '}' * ob
        try:
            return json.loads(truncated)
        except json.JSONDecodeError:
            pass
    text = re.sub(r',\s*}', '}', text)
    text = re.sub(r',\s*]', ']', text)
    return json.loads(text)


def read_pdf_file(uploaded_file):
    if not PDF_SUPPORT:
        return "PDF not supported", 0
    try:
        pdf_bytes = uploaded_file.read()
        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = ""
        pages = len(pdf_doc)
        for i in range(pages):
            text += f"\n--- Page {i+1} ---\n{pdf_doc[i].get_text()}"
        pdf_doc.close()
        return text[:8000], pages
    except Exception as e:
        return f"ERROR: {str(e)}", 0


def read_docx_file(uploaded_file):
    try:
        doc = Document(io.BytesIO(uploaded_file.read()))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return text[:8000], len(doc.paragraphs)
    except Exception as e:
        return f"ERROR: {str(e)}", 0


def generate_docx(ra_data, project_name, topic):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RISK ASSESSMENT")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Calibri"

    # Info block
    t = doc.add_table(rows=4, cols=2)
    t.style = "Table Grid"
    info_rows = [
        ("Entity Name:", "Exeed National General Contracting LLC (ENGC)"),
        ("Project:", project_name),
        ("Activity:", topic),
        ("Date:", date.today().strftime("%d-%b-%Y")),
    ]
    for i, (lbl, val) in enumerate(info_rows):
        t.rows[i].cells[0].text = lbl
        t.rows[i].cells[1].text = val
        shade_cell(t.rows[i].cells[0], "C0C0C0")
        for cell in t.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"

    doc.add_paragraph()

    # ── MAIN RA TABLE — 14 COLUMNS ──
    # CORRECT COLUMN ORDER:
    # 0=S/N | 1=Activity Element | 2=Hazards | 3=Who Harmed | 4=Consequences
    # 5=P | 6=S | 7=R | 8=Initial Risk | 9=Controls
    # 10=P | 11=S | 12=R | 13=Residual Risk

    tbl = doc.add_table(rows=2, cols=14)
    tbl.style = "Table Grid"

    # Header row 1
    h1 = tbl.rows[0]
    h1.cells[5].merge(h1.cells[7])
    h1.cells[10].merge(h1.cells[12])

    # CORRECTED HEADER ORDER
    header1 = {
        0: "S/N",
        1: "Activity Element",
        2: "Hazards / Impact",
        3: "Who Might Be Harmed and How?",      # ← COLUMN 3 (before consequences)
        4: "Risk & Potential Consequences",       # ← COLUMN 4 (after who harmed)
        5: "Risk Classification",
        8: "Initial Risk Level",
        9: "Controls",
        10: "Revised Risk Classification",
        13: "Residual Risk Level"
    }
    for idx, txt in header1.items():
        c = h1.cells[idx]
        c.text = txt
        shade_cell(c, "C0C0C0")
        for para in c.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.name = "Calibri"
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Header row 2 - P S R subheaders
    h2 = tbl.rows[1]
    sub = {5: "P", 6: "S", 7: "R\nPxS", 10: "P", 11: "S", 12: "R\nPxS"}
    for idx in range(14):
        c = h2.cells[idx]
        c.text = sub.get(idx, "")
        shade_cell(c, "BFBFBF")
        for para in c.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.name = "Calibri"
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows — CORRECTED COLUMN MAPPING
    for item in ra_data.get("rows", []):
        row = tbl.add_row()
        vals = {
            0: str(item.get("sn", "")),
            1: str(item.get("activity_element", "")),
            2: str(item.get("hazards", "")),
            3: str(item.get("who_harmed", "")),        # ← Column 3 = Who Harmed
            4: str(item.get("consequences", "")),       # ← Column 4 = Consequences
            5: str(item.get("prob_initial", "")),
            6: str(item.get("sev_initial", "")),
            7: str(item.get("risk_initial", "")),
            8: str(item.get("risk_level_initial", "")),
            9: str(item.get("controls", "")),
            10: str(item.get("prob_residual", "")),
            11: str(item.get("sev_residual", "")),
            12: str(item.get("risk_residual", "")),
            13: str(item.get("risk_level_residual", "")),
        }
        for idx, val in vals.items():
            row.cells[idx].text = val
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in row.cells[idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
                    run.font.name = "Calibri"

        shade_cell(row.cells[8],
                   get_risk_color(item.get("risk_level_initial", "LOW")))
        shade_cell(row.cells[13],
                   get_risk_color(item.get("risk_level_residual", "LOW")))

    # Legal references row
    if ra_data.get("legal_references"):
        ref_row = tbl.add_row()
        ref_row.cells[0].merge(ref_row.cells[13])
        ref_row.cells[0].text = "Legal References: " + str(ra_data.get("legal_references", ""))
        shade_cell(ref_row.cells[0], "D9E1F2")
        for para in ref_row.cells[0].paragraphs:
            for run in para.runs:
                run.font.size = Pt(8)
                run.font.name = "Calibri"
                run.bold = True

    # Signature block
    doc.add_paragraph()
    sig = doc.add_table(rows=2, cols=3)
    sig.style = "Table Grid"
    for i, lbl in enumerate(["Prepared By:", "Reviewed By:", "Approved By:"]):
        sig.rows[0].cells[i].text = lbl
        shade_cell(sig.rows[0].cells[i], "C0C0C0")
        for para in sig.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Calibri"
    for i, d in enumerate(["HSE Engineer", "HSE Manager", "Project Manager"]):
        sig.rows[1].cells[i].text = (
            "Name: _______________\n"
            "Designation: " + d + "\n"
            "Date: ____________\n"
            "Signature: ________"
        )
        for para in sig.rows[1].cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(8)
                run.font.name = "Calibri"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_prompt(activity, project):
    return (
        "Generate Risk Assessment for the following construction activity.\n"
        "Activity: " + activity + "\n"
        "Project: " + project + "\n"
        "Location: Abu Dhabi UAE\n"
        "Season: Summer\n"
        "INSTRUCTIONS:\n"
        "1. Break the activity into exactly 10 specific construction steps\n"
        "2. Every row must be unique and specific to this activity\n"
        "3. No generic rows - only steps that belong to this specific activity\n"
        "4. Row 9 must be Heat Stress - CoP 11.0 / MOHRE Resolution 44/2022\n"
        "5. Row 10 must be Emergency Response and First Aid - CoP 4.0\n"
        "6. Keep all text fields under 50 words each\n"
        "7. Include specific ADOSH CoP numbers in controls and legal references\n"
        "8. JSON column order: sn, activity_element, hazards, who_harmed, consequences, then risk scores"
    )


def build_file_prompt(activity, project, file_text, extra):
    prompt = (
        "Analyze this document and generate a Risk Assessment.\n"
        "Activity: " + activity + "\n"
        "Project: " + project + "\n"
        "Location: Abu Dhabi UAE\n"
        "Season: Summer\n"
        "INSTRUCTIONS:\n"
        "1. Extract all activities from the document\n"
        "2. Generate 10 specific rows based on document content\n"
        "3. Row 9 must be Heat Stress - CoP 11.0\n"
        "4. Row 10 must be Emergency Response - CoP 4.0\n"
        "5. Keep all text fields under 50 words\n"
        "6. Include specific ADOSH CoP numbers in all controls\n"
        "7. JSON column order: sn, activity_element, hazards, who_harmed, consequences, then risk scores\n"
        "Document content:\n" + file_text
    )
    if extra:
        prompt += "\nAdditional instructions: " + extra
    return prompt


def call_api_and_show(prompt, project_name, topic):
    try:
        client = anthropic.Anthropic(
            api_key=os.environ.get("ANTHROPIC_API_KEY")
        )
        msg = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=8192,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt}]
        )
        raw = msg.content[0].text

        with st.expander("🔍 Debug - Raw API Response"):
            st.text(raw[:500])

        ra_data = extract_json(raw)

        if "rows" not in ra_data or len(ra_data["rows"]) == 0:
            st.error("No rows generated. Please try again.")
            return

        st.success("✅ Generated " + str(len(ra_data["rows"])) + " activity rows!")

        preview = []
        for r in ra_data["rows"]:
            preview.append({
                "S/N": r.get("sn", ""),
                "Activity": str(r.get("activity_element", ""))[:35],
                "Who Harmed": str(r.get("who_harmed", ""))[:30] + "...",
                "Hazard": str(r.get("hazards", ""))[:25] + "...",
                "Initial": str(r.get("risk_level_initial", "")) + "(" + str(r.get("risk_initial", "")) + ")",
                "Residual": str(r.get("risk_level_residual", "")) + "(" + str(r.get("risk_residual", "")) + ")"
            })
        st.dataframe(preview, use_container_width=True)

        try:
            buf = generate_docx(ra_data, project_name, topic)
            fname = "ENGC_RA_" + topic.replace(" ", "_") + "_" + date.today().strftime("%d%b%Y") + ".docx"
            st.download_button(
                label="⬇️ Download Risk Assessment (DOCX)",
                data=buf,
                file_name=fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            st.success("✅ Click the button above to download your DOCX!")
        except Exception as de:
            st.error("DOCX error: " + str(de))
            st.download_button(
                "⬇️ Download JSON (backup)",
                data=json.dumps(ra_data, indent=2),
                file_name="RA_" + topic + ".json",
                mime="application/json",
                use_container_width=True
            )

        if ra_data.get("legal_references"):
            with st.expander("📚 ADOSH CoP Legal References"):
                st.info(ra_data["legal_references"])

    except json.JSONDecodeError as je:
        st.error("JSON parse error: " + str(je))
        st.info("Please click Generate again.")
    except Exception as e:
        st.error("Error: " + str(e))
        if "api_key" in str(e).lower() or "auth" in str(e).lower():
            st.warning("Check your API key in Streamlit Secrets!")


# ── SIDEBAR ──
with st.sidebar:
    st.markdown("### ⚙️ Settings")
    project_name = st.selectbox("Select Project:", [
        "Bloom Living – Almeria (Residential Villas)",
        "Bloom Living – Cordoba Community Center",
        "Bloom Living – Toledo Community Center",
        "Bloom Living Service Center (Plot C13, Zayed City)",
    ])
    st.markdown("---")
    st.markdown("**📋 Quick Topics:**")
    for topic in [
        "Fire Prevention", "Working at Heights", "Excavation Works",
        "Crane Operations", "Electrical Works", "Confined Space Entry",
        "Heat Stress Management", "Night Shift Work", "Chemical Handling",
        "Scaffolding Erection", "Concrete Works", "Lifting Operations"
    ]:
        if st.button("📄 " + topic, key="q_" + topic,
                     use_container_width=True):
            st.session_state.quick_input = topic

    st.markdown("---")
    st.markdown("**📖 ADOSH CoP Quick Reference:**")
    with st.expander("View CoP List"):
        st.markdown("""
        - **CoP 2.0** — PPE
        - **CoP 4.0** — First Aid
        - **CoP 11.0** — Safety in the Heat
        - **CoP 15.0** — Electrical Safety
        - **CoP 21.0** — Permit to Work
        - **CoP 23.0** — Working at Heights
        - **CoP 26.0** — Scaffolding
        - **CoP 27.0** — Confined Spaces
        - **CoP 28.0** — Hot Work
        - **CoP 29.0** — Excavation
        - **CoP 34.0** — Lifting Equipment
        - **CoP 36.0** — Plant & Equipment
        - **CoP 53.1** — OSH Construction Plan
        """)

# ── TABS ──
tab1, tab2, tab3 = st.tabs([
    "✏️ Type a Topic",
    "📎 Upload PDF / DOCX",
    "💬 HSE Chatbot"
])

# ── TAB 1 ──
with tab1:
    st.info("Type any construction activity → Click Generate → Download DOCX")
    user_input = st.text_input(
        "🔍 Enter Activity / Topic:",
        value=st.session_state.get("quick_input", ""),
        placeholder="e.g. Excavation Works, Scaffolding Erection..."
    )
    if st.button("🚀 Generate Risk Assessment", type="primary",
                 use_container_width=True, key="g1"):
        if not user_input.strip():
            st.warning("Please enter a topic first.")
        else:
            with st.spinner("⏳ Generating RA for: " + user_input + "..."):
                call_api_and_show(
                    build_prompt(user_input, project_name),
                    project_name,
                    user_input
                )

# ── TAB 2 ──
with tab2:
    st.info("Upload PDF or DOCX → Bot reads it → Generates RA")
    uf = st.file_uploader("Choose file:", type=["pdf", "docx", "doc"])
    if uf:
        st.success("✅ " + uf.name + " (" + str(round(len(uf.getvalue())/1024, 1)) + " KB)")
        extra = st.text_area("Extra instructions:", height=60)
        if st.button("🚀 Generate RA from File", type="primary",
                     use_container_width=True, key="g2"):
            with st.spinner("Reading file..."):
                uf.seek(0)
                if uf.name.lower().endswith(".pdf"):
                    txt, pg = read_pdf_file(uf)
                else:
                    txt, pg = read_docx_file(uf)
                if "ERROR" in str(txt):
                    st.error(txt)
                else:
                    tname = uf.name.rsplit(".", 1)[0].replace("_", " ")
                    call_api_and_show(
                        build_file_prompt(tname, project_name, txt, extra),
                        project_name,
                        tname
                    )

# ── TAB 3 — HSE CHATBOT ──
with tab3:
    st.markdown("### 💬 HSE Assistant — Ask Anything About ADOSH")
    st.info("Ask any HSE question. Answers reference specific ADOSH CoP numbers.")

    if "chat_messages" not in st.session_state:
        st.session_state.chat_messages = []

    st.markdown("**⚡ Quick Questions:**")
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("🌡️ Heat stress CoP 11.0", key="cq1", use_container_width=True):
            st.session_state.chat_input_val = "What are the heat stress requirements under ADOSH CoP 11.0 and MOHRE Resolution 44/2022?"
        if st.button("📋 Permit to Work CoP 21.0", key="cq2", use_container_width=True):
            st.session_state.chat_input_val = "What are the Permit to Work requirements under ADOSH CoP 21.0?"
    with col2:
        if st.button("🏗️ Working at Heights CoP 23.0", key="cq3", use_container_width=True):
            st.session_state.chat_input_val = "What are the working at heights requirements under ADOSH CoP 23.0?"
        if st.button("🕳️ Confined Space CoP 27.0", key="cq4", use_container_width=True):
            st.session_state.chat_input_val = "What are confined space entry requirements under ADOSH CoP 27.0?"
    with col3:
        if st.button("🏋️ Lifting CoP 34.0", key="cq5", use_container_width=True):
            st.session_state.chat_input_val = "What are lifting equipment requirements under ADOSH CoP 34.0?"
        if st.button("⚡ Electrical CoP 15.0", key="cq6", use_container_width=True):
            st.session_state.chat_input_val = "What are electrical safety requirements under ADOSH CoP 15.0?"

    st.markdown("---")

    for msg in st.session_state.chat_messages:
        if msg["role"] == "user":
            st.markdown(
                "<div style='background-color:#1e3a5f;padding:10px;"
                "border-radius:8px;margin:5px 0;color:white;'>"
                "👤 <b>You:</b> " + msg["content"] + "</div>",
                unsafe_allow_html=True
            )
        else:
            st.markdown(
                "<div style='background-color:#1a4a2e;padding:10px;"
                "border-radius:8px;margin:5px 0;color:white;'>"
                "🦺 <b>HSE Assistant:</b><br>" + msg["content"] + "</div>",
                unsafe_allow_html=True
            )

    st.markdown("---")
    chat_input = st.text_area(
        "💬 Type your HSE question:",
        value=st.session_state.get("chat_input_val", ""),
        placeholder="e.g. What is ADOSH CoP 23.0 requirement for edge protection?",
        height=80,
        key="chat_box"
    )

    col_send, col_clear = st.columns([3, 1])
    with col_send:
        if st.button("📤 Send Question", type="primary",
                     use_container_width=True, key="send_chat"):
            question = chat_input.strip()
            if not question:
                st.warning("Please type a question first.")
            else:
                st.session_state.chat_messages.append({
                    "role": "user", "content": question
                })
                st.session_state.chat_input_val = ""
                with st.spinner("🤔 Checking ADOSH regulations..."):
                    try:
                        client = anthropic.Anthropic(
                            api_key=os.environ.get("ANTHROPIC_API_KEY")
                        )
                        messages = [
                            {"role": m["role"], "content": m["content"]}
                            for m in st.session_state.chat_messages
                        ]
                        response = client.messages.create(
                            model="claude-sonnet-4-6",
                            max_tokens=1500,
                            system=CHAT_SYSTEM_PROMPT,
                            messages=messages
                        )
                        answer = response.content[0].text
                        st.session_state.chat_messages.append({
                            "role": "assistant", "content": answer
                        })
                        st.rerun()
                    except Exception as e:
                        st.error("Error: " + str(e))

    with col_clear:
        if st.button("🗑️ Clear", use_container_width=True, key="clear_chat"):
            st.session_state.chat_messages = []
            st.session_state.chat_input_val = ""
            st.rerun()

    if st.session_state.chat_messages:
        st.caption("💬 " + str(len(st.session_state.chat_messages) // 2) + " questions this session.")

st.markdown("---")
st.caption(
    "⚠️ AI output must be reviewed by a competent HSE professional. "
    "| ADOSH-SF Version 4.0 | © ENGC 2026"
)
